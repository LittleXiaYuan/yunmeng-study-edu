#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from lxml import etree

def insert_bold_paragraph_after(paragraph, text, font_size=24):
    nsmap = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    new_p = etree.Element(nsmap + "p")
    # 构建 r
    r = etree.SubElement(new_p, nsmap + "r")
    rPr = etree.SubElement(r, nsmap + "rPr")
    rFonts = etree.SubElement(rPr, nsmap + "rFonts")
    rFonts.set(nsmap + "eastAsia", "HarmonyOS Sans SC")
    sz = etree.SubElement(rPr, nsmap + "sz")
    sz.set(nsmap + "val", str(font_size * 2))
    szCs = etree.SubElement(rPr, nsmap + "szCs")
    szCs.set(nsmap + "val", str(font_size * 2))
    t = etree.SubElement(r, nsmap + "t")
    t.text = text
    paragraph._element.addnext(new_p)


def insert_bold_paragraph(doc, text, font_size=24):
    p = doc.add_paragraph(text)
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(font_size)


def main():
    template_path = r"C:\Code\AI\Study\docs\competition\2026中国高校计算机大赛人工智能创意赛初赛（鸿蒙赛道）作品说明文档模板.docx"
    doc = Document(template_path)
    print(f"模板共 {len(doc.paragraphs)} 段落")

    target_indices = {
        "team_info": 80,
        "one_liner": 87,
        "design": 93,
        "introduction": 100
    }

    for i, para in enumerate(doc.paragraphs):
        if i == target_indices["team_info"]:
            doc.add_paragraph()
            doc.add_paragraph("【团队信息】")
            doc.add_paragraph("参赛队伍：云元队")
            doc.add_paragraph("指导教师：[待填写学校教师姓名]")
            doc.add_paragraph("参赛学校：[待填写学校名称]")
            doc.add_paragraph()
            print(f"✓ 团队信息插入位置: 第 {i+1} 段")
            break

    for i, para in enumerate(doc.paragraphs):
        if i == target_indices["one_liner"]:
            insert_bold_paragraph_after(para, "AI Agent驱动的沉浸式课堂智能伴学与教学辅助系统", 24)
            print(f"✓ 一句话创意插入位置: 第 {i+1} 段")
            break

    for i, para in enumerate(doc.paragraphs):
        if i == target_indices["design"]:
            insert_bold_paragraph(doc, "七、技术方案与架构", 24)
            insert_bold_paragraph_after(para, "图 1 系统整体架构（云元服务架构）", 24)
            insert_bold_paragraph_after(para, "说明：架构图可放项目根目录 architecture.png，调整宽度约 5 英寸")
            doc.add_paragraph()
            print(f"✓ 技术方案与架构插入位置: 第 {i+1} 段")
            break

    intro_text = (
        "“云元——沉浸式AI伴学系统”以鸿蒙全场景 AI Agent 为核心切入点，针对当前课堂数字化教学中的核心痛点提供沉浸式伴学解决方案。"
        "教学场景中，学生基础差异大、教师精力有限，且通用大模型容易坍缩为“答案机”，削弱学生的自主思考过程；"
        "同时教师零散资料难以系统化沉淀，知识库检索效率较低，学生学习过程缺乏个性化正向反馈与分层引导。"
        "\n\n"
        "为解决上述痛点，本作品在教师端设计“智能导入 Agent”：支持挂载 zip/docx/pdf 等多格式教学资料，"
        "通过意图识别自动拆解教案、练习与噪声，仅将知识价值高的内容过滤入库，并返回导入报告卡供教师确认；"
        "学生对“文件上传 → RAG 知识库 → 信任分门控 → page_context 全场景教练”链路形成训练数据闭环，增强教师的掌控感与系统的可信度。"
        "\n\n"
        "在学生端，以“学生教练”为核心提供个性化伴学：依据作业题干、阶段选择、答题进度实时组装 page_context，"
        "结合信任分决定只提问、给提示还是讲思路，从 locked → hint → partial → explain 四级渐进式引导；"
        "同时引用老师上传的课程知识库片段给出具体章节引用，确保引导有据可循。"
        "多 Agent（Teacher 解析、Tutor 伴学、Evaluator 评价、Reflector 汇总共性）协同驱动教学闭环，"
        "本地优先的伴学策略可在鸿蒙小艺 Agent、PC 前端、穿戴与车机中复用。"
        "\n\n"
        "系统能力层面，构建“学校 → 班级 → 课程 → 用户”四层组织模型，超管可对多校多班级进行数据隔离与安全管控；"
        "教师和学生在各自授权范围内对文章、测试、会话等进行细粒度访问与导出；"
        "本地先行策略结合适配层支持任意 LLM 切换。在鸿蒙全场景上，小艺作为 Agent 入口提供自然对话与主动提醒，"
        "笔记与作业图文自动流转至教室，实现多终端教学协同落地。"
        "\n\n"
        "本项目已在技术实现与测试层面完成，涵盖用户隔离、权限矩阵、多 Agent 协作与鸿蒙 Agent 入口，"
        "具有明确的应用推广路径与商业化潜力。"
    )

    for i, para in enumerate(doc.paragraphs):
        if i == target_indices["introduction"]:
            insert_bold_paragraph(doc, "八、作品介绍", 24)
            insert_bold_paragraph_after(para, intro_text, 21)
            print(f"✓ 作品介绍插入位置: 第 {i+1} 段")
            break

    output_path = r"C:\Code\AI\Study\dist-c4\01-作品说明文档-云元队.docx"
    doc.save(output_path)
    print(f"\n✓ 最终文档已保存: {output_path}")


if __name__ == "__main__":
    main()
