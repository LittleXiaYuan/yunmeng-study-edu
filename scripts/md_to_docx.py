"""
md → docx 转换器
简单但够用：标题 / 段落 / 代码块 / 列表 / 粗体 / 表格 / 引用
用于课程作业交付：把 Markdown 转成 .docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


def add_runs_with_inline_bold(paragraph, text):
    """处理 **粗体** 标记"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def parse_table(doc, lines, start_idx):
    """解析 Markdown 表格，返回 (table, next_idx)"""
    header = [c.strip() for c in lines[start_idx].strip("|").split("|")]
    # 跳过 |---|---| 分隔行
    idx = start_idx + 2
    rows = []
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        cells = [c.strip() for c in lines[idx].strip("|").split("|")]
        rows.append(cells)
        idx += 1
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    # header
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
    # body
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i + 1].cells):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs_with_inline_bold(p, cell_text)
    return table, idx


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    # 中文字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")
    i = 0
    in_code = False
    code_buf = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:].strip()
                code_buf = []
            else:
                in_code = False
                # 写代码块
                if code_buf:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    run = p.add_run("\n".join(code_buf))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9.5)
                    if code_lang:
                        # 在代码块前加个小标签
                        pass
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            add_runs_with_inline_bold(h, heading_text)
            i += 1
            continue

        # 表格
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
            _, i = parse_table(doc, lines, i)
            continue

        # 水平线
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # 引用
        m = re.match(r"^>\s*(.*)$", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run("▎ " + m.group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 无序列表
        m = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if m:
            indent = len(m.group(1))
            content = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3 + indent * 0.2)
            add_runs_with_inline_bold(p, content)
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if m:
            indent = len(m.group(1))
            content = m.group(3)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.3 + indent * 0.2)
            add_runs_with_inline_bold(p, content)
            i += 1
            continue

        # 空行
        if line.strip() == "":
            i += 1
            continue

        # 普通段落（合并连续行）
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != "" and not re.match(
            r"^(#{1,6}\s|[-*]\s|\d+\.\s|>\s|```|---|\|)", lines[i]
        ):
            para_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        add_runs_with_inline_bold(p, " ".join(para_lines))

    doc.save(docx_path)
    print(f"  ✓ {md_path.name} → {docx_path.name}  ({docx_path.stat().st_size//1024} KB)")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("usage: python md_to_docx.py <src_dir> [dst_dir]")
        sys.exit(1)
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src
    dst.mkdir(parents=True, exist_ok=True)
    for md in sorted(src.glob("*.md")):
        docx = dst / (md.stem + ".docx")
        try:
            md_to_docx(md, docx)
        except Exception as e:
            print(f"  ✗ {md.name}: {e}")
